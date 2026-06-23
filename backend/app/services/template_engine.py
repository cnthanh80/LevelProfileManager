from __future__ import annotations

import hashlib
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import HTTPException
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from app.services.pdf_font import apply_unicode_styles
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.document_template import DocumentTemplate
from app.models.evidence_document import EvidenceDocument
from app.models.exported_document import ExportedDocument
from app.models.information_system import InformationSystem
from app.models.level_profile import LevelProfile
from app.models.profile_requirement_answer import ProfileRequirementAnswer
from app.models.security_requirement import SecurityRequirement


EXPORT_ROOT = Path("/app/storage/exports")

GOVERNMENT_DOCUMENT_TYPES = {
    "PROFILE_EXPLANATION": "Thuyết minh hồ sơ đề xuất cấp độ",
    "CONSULTATION_REQUEST": "Công văn xin ý kiến chuyên môn",
    "APPROVAL_SUBMISSION": "Tờ trình phê duyệt hồ sơ đề xuất cấp độ",
    "APPROVAL_DECISION": "Quyết định phê duyệt cấp độ an toàn hệ thống thông tin",
    "CHECKLIST_APPENDIX": "Phụ lục checklist đáp ứng yêu cầu ATTT",
    "SYSTEM_SUMMARY_REPORT": "Báo cáo tổng hợp danh sách hệ thống theo cấp độ",
    "OVERDUE_REVIEW_REPORT": "Báo cáo hệ thống quá hạn rà soát",
}

DEFAULT_TEMPLATES = [
    ("TPL_PROFILE_EXPLANATION_V1", "Mẫu thuyết minh hồ sơ đề xuất cấp độ", "PROFILE_EXPLANATION", 10),
    ("TPL_CONSULTATION_REQUEST_V1", "Mẫu công văn xin ý kiến chuyên môn", "CONSULTATION_REQUEST", 20),
    ("TPL_APPROVAL_SUBMISSION_V1", "Mẫu tờ trình phê duyệt", "APPROVAL_SUBMISSION", 30),
    ("TPL_APPROVAL_DECISION_V1", "Mẫu quyết định phê duyệt", "APPROVAL_DECISION", 40),
    ("TPL_CHECKLIST_APPENDIX_V1", "Mẫu phụ lục checklist", "CHECKLIST_APPENDIX", 50),
]


def seed_default_templates(db: Session) -> int:
    count = 0
    for code, name, document_type, sort_order in DEFAULT_TEMPLATES:
        item = db.scalar(select(DocumentTemplate).where(DocumentTemplate.code == code))
        if not item:
            db.add(DocumentTemplate(
                code=code,
                name=name,
                document_type=document_type,
                version="1.0",
                description="Mẫu mặc định sinh tự động, có thể thay thế bằng template DOCX tùy biến ở phase production.",
                agency_name="NGÂN HÀNG CHÍNH SÁCH XÃ HỘI",
                file_format="docx",
                is_active=True,
                sort_order=sort_order,
            ))
            count += 1
    db.commit()
    return count


def _safe(value) -> str:
    return "" if value is None else str(value)


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _get_profile(db: Session, profile_id: int) -> tuple[LevelProfile, InformationSystem]:
    profile = db.get(LevelProfile, profile_id)
    if not profile:
        raise HTTPException(status_code=404, detail="Level profile not found")
    system = db.get(InformationSystem, profile.information_system_id)
    if not system:
        raise HTTPException(status_code=404, detail="Information system not found")
    return profile, system


def _checklist_rows(db: Session, profile_id: int):
    stmt = (
        select(ProfileRequirementAnswer, SecurityRequirement)
        .join(SecurityRequirement, ProfileRequirementAnswer.requirement_id == SecurityRequirement.id)
        .where(ProfileRequirementAnswer.profile_id == profile_id)
        .order_by(SecurityRequirement.group_name, SecurityRequirement.category, SecurityRequirement.code)
    )
    return db.execute(stmt).all()


def _evidence_count(db: Session, profile_id: int) -> int:
    return len(db.scalars(select(EvidenceDocument).where(EvidenceDocument.profile_id == profile_id)).all())


def _add_center(doc: Document, text: str, bold: bool = False, size: int = 12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def _add_right_italic(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def _set_doc_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Nội dung"
    table.rows[0].cells[1].text = "Thông tin"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    return table


def _add_formal_header(doc: Document, agency_name: str | None):
    table = doc.add_table(rows=1, cols=2)
    left, right = table.rows[0].cells
    left.text = f"{agency_name or 'CƠ QUAN, TỔ CHỨC'}\nĐƠN VỊ QUẢN LÝ HỆ THỐNG"
    right.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
    doc.add_paragraph("")


def _add_signature_block(doc: Document, signer_title: str | None, signer_name: str | None):
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Nơi nhận:\n- Như trên;\n- Lưu: VT, CNTT."
    sig = table.rows[0].cells[1]
    sig.text = f"{signer_title or 'THỦ TRƯỞNG ĐƠN VỊ'}\n\n\n\n{signer_name or '(Ký, ghi rõ họ tên)'}"
    for p in sig.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True


def _build_profile_explanation(doc: Document, db: Session, profile: LevelProfile, system: InformationSystem):
    doc.add_heading("I. THÔNG TIN CHUNG VỀ HỆ THỐNG THÔNG TIN", level=1)
    _add_key_value_table(doc, [
        ("Tên hệ thống", _safe(system.name)),
        ("Mã hệ thống", _safe(system.code)),
        ("Mục tiêu", _safe(system.purpose)),
        ("Phạm vi", _safe(system.scope)),
        ("Chức năng chính", _safe(system.main_functions)),
        ("Nhóm người dùng", _safe(system.user_groups)),
        ("Loại dữ liệu xử lý", _safe(system.data_types)),
        ("Mô hình triển khai", _safe(system.deployment_model)),
        ("Môi trường", _safe(system.environment)),
    ])
    doc.add_heading("II. CĂN CỨ ĐỀ XUẤT CẤP ĐỘ", level=1)
    doc.add_paragraph(_safe(profile.basis_for_level) or f"Hồ sơ đề xuất cấp độ {profile.proposed_level} căn cứ mức độ ảnh hưởng khi mất an toàn thông tin.")
    doc.add_heading("III. THUYẾT MINH PHẠM VI VÀ KIẾN TRÚC", level=1)
    doc.add_paragraph("Phạm vi hệ thống: " + (_safe(profile.system_scope_description) or "Chưa cập nhật."))
    doc.add_paragraph("Kiến trúc kỹ thuật: " + (_safe(profile.technical_architecture) or "Chưa cập nhật."))
    doc.add_heading("IV. ĐÁNH GIÁ TÁC ĐỘNG", level=1)
    _add_key_value_table(doc, [
        ("Mất tính bí mật", _safe(profile.confidentiality_impact)),
        ("Mất tính toàn vẹn", _safe(profile.integrity_impact)),
        ("Mất tính sẵn sàng", _safe(profile.availability_impact)),
    ])
    doc.add_heading("V. TỔNG HỢP MINH CHỨNG", level=1)
    doc.add_paragraph(f"Số tài liệu minh chứng đã gắn với hồ sơ: {_evidence_count(db, profile.id)}")


def _build_checklist_appendix(doc: Document, db: Session, profile: LevelProfile):
    doc.add_heading("PHỤ LỤC CHECKLIST ĐÁP ỨNG YÊU CẦU ATTT", level=1)
    rows = _checklist_rows(db, profile.id)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["STT", "Mã", "Nhóm", "Danh mục", "Yêu cầu", "Trạng thái", "Minh chứng"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for idx, (answer, req) in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(idx)
        cells[1].text = _safe(req.code)
        cells[2].text = _safe(req.group_name)
        cells[3].text = _safe(req.category)
        cells[4].text = _safe(req.title)
        cells[5].text = _safe(answer.status)
        cells[6].text = str(answer.evidence_count or 0)


def _build_submission(doc: Document, profile: LevelProfile, system: InformationSystem):
    doc.add_paragraph("Kính gửi: Lãnh đạo đơn vị")
    doc.add_paragraph(
        f"Căn cứ nhu cầu bảo đảm an toàn hệ thống thông tin theo cấp độ, đơn vị trình phê duyệt hồ sơ đề xuất cấp độ đối với hệ thống {system.name}."
    )
    _add_key_value_table(doc, [
        ("Mã hồ sơ", profile.profile_code),
        ("Tên hệ thống", system.name),
        ("Cấp độ đề xuất", str(profile.proposed_level)),
        ("Trạng thái hồ sơ", profile.status),
    ])
    doc.add_paragraph("Kính đề nghị Lãnh đạo xem xét, phê duyệt để làm căn cứ triển khai các biện pháp bảo đảm an toàn thông tin phù hợp.")


def _build_consultation_request(doc: Document, profile: LevelProfile, system: InformationSystem):
    doc.add_paragraph("Kính gửi: Đơn vị chuyên trách về an toàn thông tin")
    doc.add_paragraph(
        f"Đơn vị gửi hồ sơ đề xuất cấp độ an toàn hệ thống thông tin đối với {system.name} để xin ý kiến chuyên môn trước khi hoàn thiện trình phê duyệt."
    )
    _add_key_value_table(doc, [
        ("Mã hệ thống", system.code),
        ("Mã hồ sơ", profile.profile_code),
        ("Cấp độ đề xuất", str(profile.proposed_level)),
        ("Mô hình triển khai", _safe(system.deployment_model)),
    ])
    doc.add_paragraph("Kính đề nghị Quý đơn vị xem xét, cho ý kiến đối với hồ sơ kèm theo.")


def _build_approval_decision(doc: Document, profile: LevelProfile, system: InformationSystem):
    doc.add_paragraph("QUYẾT ĐỊNH")
    doc.add_paragraph(f"Về việc phê duyệt cấp độ an toàn hệ thống thông tin đối với {system.name}")
    doc.add_paragraph("Điều 1. Phê duyệt cấp độ an toàn hệ thống thông tin như sau:")
    _add_key_value_table(doc, [
        ("Tên hệ thống", system.name),
        ("Mã hệ thống", system.code),
        ("Cấp độ được phê duyệt", str(profile.proposed_level)),
        ("Đơn vị vận hành", _safe(system.operator_org_id)),
    ])
    doc.add_paragraph("Điều 2. Đơn vị quản lý, vận hành hệ thống có trách nhiệm triển khai đầy đủ các biện pháp bảo đảm an toàn thông tin theo cấp độ được phê duyệt.")
    doc.add_paragraph("Điều 3. Quyết định này có hiệu lực kể từ ngày ký.")


def generate_government_document(
    db: Session,
    profile_id: int,
    document_type: str,
    file_format: str,
    generated_by: int | None,
    agency_name: str | None = None,
    signer_title: str | None = None,
    signer_name: str | None = None,
    place_name: str | None = None,
    template_code: str | None = None,
) -> ExportedDocument:
    document_type = document_type.upper()
    file_format = file_format.lower()
    if document_type not in GOVERNMENT_DOCUMENT_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported document_type. Allowed: {', '.join(GOVERNMENT_DOCUMENT_TYPES)}")
    if file_format not in {"docx", "pdf"}:
        raise HTTPException(status_code=400, detail="Unsupported file_format. Allowed: docx, pdf")
    if template_code:
        template = db.scalar(select(DocumentTemplate).where(DocumentTemplate.code == template_code, DocumentTemplate.is_active == True))
        if not template:
            raise HTTPException(status_code=404, detail="Document template not found or inactive")
        document_type = template.document_type
        agency_name = agency_name or template.agency_name

    profile, system = _get_profile(db, profile_id)
    title = GOVERNMENT_DOCUMENT_TYPES[document_type]
    EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    stored_filename = f"govdoc_{profile_id}_{document_type.lower()}_{uuid4().hex}.{file_format}"
    output_path = EXPORT_ROOT / stored_filename

    if file_format == "docx":
        active_template = None
        if template_code:
            active_template = template if template.template_path and template.file_format.lower() == "docx" else None
        else:
            active_template = get_active_template(db, document_type)
        if active_template and active_template.template_path:
            context = build_template_context(
                db,
                profile_id,
                agency_name=agency_name or active_template.agency_name,
                place_name=place_name or "Hà Nội",
                signer_title=signer_title or "THỦ TRƯỞNG ĐƠN VỊ",
                signer_name=signer_name,
            )
            render_docx_template(active_template.template_path, output_path, context)
        else:
            doc = Document()
            _set_doc_style(doc)
            _add_formal_header(doc, agency_name)
            _add_right_italic(doc, f"{place_name or 'Hà Nội'}, ngày {datetime.now().day:02d} tháng {datetime.now().month:02d} năm {datetime.now().year}")
            _add_center(doc, title.upper(), bold=True, size=14)
            _add_center(doc, f"Hồ sơ: {profile.profile_code}", bold=False, size=12)
            if document_type == "PROFILE_EXPLANATION":
                _build_profile_explanation(doc, db, profile, system)
            elif document_type == "CHECKLIST_APPENDIX":
                _build_checklist_appendix(doc, db, profile)
            elif document_type == "APPROVAL_SUBMISSION":
                _build_submission(doc, profile, system)
            elif document_type == "CONSULTATION_REQUEST":
                _build_consultation_request(doc, profile, system)
            elif document_type == "APPROVAL_DECISION":
                _build_approval_decision(doc, profile, system)
            else:
                _build_profile_explanation(doc, db, profile, system)
            _add_signature_block(doc, signer_title, signer_name)
            doc.save(output_path)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        styles = getSampleStyleSheet()
        regular_font, bold_font = apply_unicode_styles(styles)
        story = [Paragraph(title.upper(), styles["Title"]), Paragraph(f"Hồ sơ: {profile.profile_code}", styles["Normal"]), Spacer(1, 12)]
        data = [["Thông tin", "Nội dung"], ["Tên hệ thống", _safe(system.name)], ["Mã hệ thống", _safe(system.code)], ["Cấp độ", str(profile.proposed_level)], ["Trạng thái", _safe(profile.status)], ["Minh chứng", str(_evidence_count(db, profile.id))]]
        table = Table(data, colWidths=[150, 340])
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("FONTNAME", (0, 0), (-1, -1), regular_font)]))
        story.append(table)
        story.append(Spacer(1, 12))
        story.append(Paragraph("Bản PDF được sinh từ Government Document Generator v1.6.", styles["Normal"]))
        SimpleDocTemplate(str(output_path), pagesize=A4).build(story)
        content_type = "application/pdf"

    item = ExportedDocument(
        profile_id=profile_id,
        document_type=document_type,
        file_format=file_format,
        title=title,
        original_filename=output_path.name,
        stored_filename=output_path.name,
        storage_path=str(output_path),
        content_type=content_type,
        file_size=output_path.stat().st_size,
        checksum_sha256=_sha256(output_path),
        generated_by=generated_by,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item

def get_active_template(db, document_type: str):
    return (
        db.query(DocumentTemplate)
        .filter(
            DocumentTemplate.document_type == document_type,
            DocumentTemplate.is_active == True
        )
        .order_by(DocumentTemplate.id.desc())
        .first()
    )