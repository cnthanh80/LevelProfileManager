from __future__ import annotations

import hashlib
import re
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.models.evidence_document import EvidenceDocument
from app.models.government_dossier import GovernmentDossier, GovernmentDossierFile
from app.models.information_system import InformationSystem
from app.models.level_profile import LevelProfile
from app.models.organization import Organization
from app.models.profile_requirement_answer import ProfileRequirementAnswer
from app.models.security_requirement import SecurityRequirement

DOSSIER_ROOT = Path("/app/storage/dossiers")

CONTENT_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "zip": "application/zip",
}

DOC_TYPES = [
    ("TO_TRINH", "01-ToTrinh.docx", "Tờ trình phê duyệt hồ sơ đề xuất cấp độ", 10),
    ("HO_SO_CAP_DO", "02-HoSoCapDo.docx", "Hồ sơ đề xuất cấp độ an toàn hệ thống thông tin", 20),
    ("XIN_Y_KIEN", "03-VanBanXinYKien.docx", "Văn bản xin ý kiến chuyên môn", 30),
    ("QUYET_DINH", "04-QuyetDinh.docx", "Quyết định phê duyệt cấp độ", 40),
    ("CHECKLIST", "05-Checklist.xlsx", "Phụ lục checklist đáp ứng yêu cầu ATTT", 50),
]


def _safe(value: object, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def _slug(value: object) -> str:
    text = _safe(value, "HSCD")
    text = re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("._")
    return text or "HSCD"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def _set_doc_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(13 if style_name == "Normal" else 14)
        except Exception:
            pass


def _add_center(doc: Document, text: str, bold: bool = False, size: int = 13) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = "Times New Roman"


def _add_kv_table(doc: Document, rows: list[tuple[str, object]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Nội dung"
    hdr[1].text = "Thông tin"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = _safe(key)
        cells[1].text = _safe(value, "Chưa cập nhật")


def _org_name(db: Session, org_id: int | None) -> str:
    if not org_id:
        return "Chưa cập nhật"
    org = db.get(Organization, org_id)
    return org.name if org else "Chưa cập nhật"


def _context(db: Session, profile_id: int) -> dict:
    profile = db.get(LevelProfile, profile_id)
    if not profile:
        raise ValueError("Level profile not found")
    system = db.get(InformationSystem, profile.information_system_id)
    answers = db.scalars(
        select(ProfileRequirementAnswer)
        .where(ProfileRequirementAnswer.profile_id == profile_id)
        .order_by(ProfileRequirementAnswer.id)
    ).all()
    requirement_ids = [a.requirement_id for a in answers]
    requirements = {}
    if requirement_ids:
        reqs = db.scalars(select(SecurityRequirement).where(SecurityRequirement.id.in_(requirement_ids))).all()
        requirements = {r.id: r for r in reqs}
    evidence = db.scalars(
        select(EvidenceDocument)
        .where(EvidenceDocument.profile_id == profile_id)
        .order_by(EvidenceDocument.created_at.desc(), EvidenceDocument.id.desc())
    ).all()
    compliant = sum(1 for a in answers if _safe(a.status).upper() == "COMPLIANT")
    na = sum(1 for a in answers if _safe(a.status).upper() == "NOT_APPLICABLE")
    total_effective = max(1, len(answers) - na)
    compliance_rate = round(compliant * 100 / total_effective, 2)
    return {
        "profile": profile,
        "system": system,
        "answers": answers,
        "requirements": requirements,
        "evidence": evidence,
        "compliance_rate": compliance_rate,
        "owner_org": _org_name(db, system.owner_org_id if system else None),
        "operator_org": _org_name(db, system.operator_org_id if system else None),
        "generated_at": datetime.now(),
    }


def _common_header(doc: Document, title: str) -> None:
    _add_center(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", True, 13)
    _add_center(doc, "Độc lập - Tự do - Hạnh phúc", True, 13)
    doc.add_paragraph()
    _add_center(doc, title.upper(), True, 15)
    doc.add_paragraph()


def _system_rows(ctx: dict) -> list[tuple[str, object]]:
    profile = ctx["profile"]
    system = ctx["system"]
    return [
        ("Tên hệ thống thông tin", system.name if system else "Chưa cập nhật"),
        ("Mã hệ thống", system.code if system else "Chưa cập nhật"),
        ("Mã hồ sơ", profile.profile_code),
        ("Đơn vị chủ quản", ctx["owner_org"]),
        ("Đơn vị vận hành", ctx["operator_org"]),
        ("Cấp độ đề xuất", profile.proposed_level),
        ("Trạng thái hồ sơ", profile.status),
        ("Mô hình triển khai", system.deployment_model if system else "Chưa cập nhật"),
        ("Môi trường", system.environment if system else "Chưa cập nhật"),
        ("Mức độ quan trọng", system.importance_level if system else "Chưa cập nhật"),
    ]


def _make_to_trinh(path: Path, ctx: dict) -> None:
    doc = Document()
    _set_doc_default_font(doc)
    _common_header(doc, "Tờ trình")
    profile = ctx["profile"]
    system = ctx["system"]
    doc.add_paragraph(f"Về việc phê duyệt Hồ sơ đề xuất cấp độ an toàn hệ thống thông tin đối với hệ thống {_safe(system.name if system else None, profile.profile_code)}")
    doc.add_paragraph("Kính trình: Lãnh đạo đơn vị")
    _add_heading(doc, "I. Căn cứ pháp lý", 1)
    for item in [
        "Nghị định số 85/2016/NĐ-CP về bảo đảm an toàn hệ thống thông tin theo cấp độ.",
        "Thông tư số 12/2022/TT-BTTTT hướng dẫn Nghị định số 85/2016/NĐ-CP.",
        "Nhu cầu quản lý, vận hành và bảo đảm an toàn thông tin của cơ quan/tổ chức.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    _add_heading(doc, "II. Thông tin hệ thống", 1)
    _add_kv_table(doc, _system_rows(ctx))
    _add_heading(doc, "III. Kết quả đề xuất cấp độ", 1)
    doc.add_paragraph(f"Căn cứ phạm vi, dữ liệu xử lý, kết nối, tác động khi mất bí mật, toàn vẹn, sẵn sàng, hồ sơ đề xuất cấp độ {profile.proposed_level} cho hệ thống.")
    doc.add_paragraph(f"Tỷ lệ đáp ứng checklist hiện tại: {ctx['compliance_rate']}%.")
    _add_heading(doc, "IV. Kiến nghị", 1)
    doc.add_paragraph("Kính đề nghị Lãnh đạo xem xét, phê duyệt hồ sơ đề xuất cấp độ để làm căn cứ triển khai, duy trì và kiểm tra các biện pháp bảo đảm an toàn thông tin.")
    doc.add_paragraph("Nơi nhận:\n- Như trên;\n- Đơn vị vận hành;\n- Lưu: Hồ sơ ATTT.")
    doc.save(path)


def _make_hoso(path: Path, ctx: dict) -> None:
    doc = Document()
    _set_doc_default_font(doc)
    _common_header(doc, "Hồ sơ đề xuất cấp độ an toàn hệ thống thông tin")
    profile = ctx["profile"]
    system = ctx["system"]
    _add_heading(doc, "Chương I. Thông tin chung", 1)
    _add_kv_table(doc, _system_rows(ctx))
    _add_heading(doc, "Chương II. Mô tả hệ thống", 1)
    doc.add_paragraph("Mục tiêu hệ thống: " + _safe(system.purpose if system else None, "Chưa cập nhật"))
    doc.add_paragraph("Phạm vi hệ thống: " + _safe(system.scope if system else profile.system_scope_description, "Chưa cập nhật"))
    doc.add_paragraph("Chức năng chính: " + _safe(system.main_functions if system else None, "Chưa cập nhật"))
    doc.add_paragraph("Nhóm người dùng: " + _safe(system.user_groups if system else None, "Chưa cập nhật"))
    doc.add_paragraph("Loại dữ liệu xử lý: " + _safe(system.data_types if system else None, "Chưa cập nhật"))
    _add_heading(doc, "Chương III. Kiến trúc kỹ thuật", 1)
    doc.add_paragraph(_safe(profile.technical_architecture, "Chưa cập nhật thuyết minh kiến trúc kỹ thuật."))
    _add_heading(doc, "Chương IV. Căn cứ xác định cấp độ", 1)
    doc.add_paragraph("Căn cứ xác định cấp độ: " + _safe(profile.basis_for_level, "Chưa cập nhật"))
    _add_kv_table(doc, [
        ("Tác động khi mất bí mật", profile.confidentiality_impact),
        ("Tác động khi mất toàn vẹn", profile.integrity_impact),
        ("Tác động khi mất sẵn sàng", profile.availability_impact),
        ("Cấp độ đề xuất", profile.proposed_level),
    ])
    _add_heading(doc, "Chương V. Hiện trạng đáp ứng yêu cầu ATTT", 1)
    doc.add_paragraph(f"Tổng số tiêu chí checklist: {len(ctx['answers'])}. Tỷ lệ đáp ứng: {ctx['compliance_rate']}%.")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(["Mã", "Yêu cầu", "Nhóm", "Trạng thái", "Minh chứng"]):
        table.rows[0].cells[i].text = h
    for answer in ctx["answers"][:80]:
        req = ctx["requirements"].get(answer.requirement_id)
        cells = table.add_row().cells
        cells[0].text = _safe(req.code if req else answer.requirement_id)
        cells[1].text = _safe(req.title if req else "Yêu cầu")
        cells[2].text = _safe(req.group_name if req else "")
        cells[3].text = _safe(answer.status)
        cells[4].text = str(answer.evidence_count or 0)
    _add_heading(doc, "Phụ lục. Danh mục tài liệu minh chứng", 1)
    for ev in ctx["evidence"][:100]:
        doc.add_paragraph(f"- {ev.title} ({ev.document_type}) - {ev.original_filename}")
    doc.save(path)


def _make_xin_y_kien(path: Path, ctx: dict) -> None:
    doc = Document()
    _set_doc_default_font(doc)
    _common_header(doc, "Văn bản xin ý kiến chuyên môn")
    profile = ctx["profile"]
    system = ctx["system"]
    doc.add_paragraph("Kính gửi: Đơn vị chuyên trách về an toàn thông tin")
    doc.add_paragraph(f"Đơn vị gửi kèm Hồ sơ đề xuất cấp độ an toàn hệ thống thông tin đối với hệ thống {_safe(system.name if system else None, profile.profile_code)} để xin ý kiến chuyên môn theo quy định.")
    _add_kv_table(doc, _system_rows(ctx))
    doc.add_paragraph("Kính đề nghị Quý đơn vị xem xét, cho ý kiến để đơn vị hoàn thiện hồ sơ trước khi trình phê duyệt chính thức.")
    doc.add_paragraph("Trân trọng.")
    doc.save(path)


def _make_quyet_dinh(path: Path, ctx: dict) -> None:
    doc = Document()
    _set_doc_default_font(doc)
    _common_header(doc, "Quyết định")
    profile = ctx["profile"]
    system = ctx["system"]
    _add_center(doc, f"Về việc phê duyệt cấp độ an toàn hệ thống thông tin đối với {_safe(system.name if system else None, profile.profile_code)}", True, 13)
    doc.add_paragraph("THỦ TRƯỞNG ĐƠN VỊ")
    for item in [
        "Căn cứ Nghị định số 85/2016/NĐ-CP về bảo đảm an toàn hệ thống thông tin theo cấp độ;",
        "Căn cứ Thông tư số 12/2022/TT-BTTTT hướng dẫn Nghị định số 85/2016/NĐ-CP;",
        "Căn cứ Hồ sơ đề xuất cấp độ an toàn hệ thống thông tin đã được rà soát;",
    ]:
        doc.add_paragraph(item)
    _add_center(doc, "QUYẾT ĐỊNH:", True, 13)
    doc.add_paragraph(f"Điều 1. Phê duyệt cấp độ {profile.proposed_level} đối với hệ thống thông tin {_safe(system.name if system else None, profile.profile_code)}.")
    doc.add_paragraph("Điều 2. Đơn vị chủ quản và đơn vị vận hành có trách nhiệm triển khai, duy trì, rà soát các biện pháp bảo đảm an toàn thông tin tương ứng với cấp độ đã được phê duyệt.")
    doc.add_paragraph("Điều 3. Quyết định này có hiệu lực kể từ ngày ký. Các đơn vị liên quan chịu trách nhiệm thi hành Quyết định này.")
    doc.add_paragraph("Nơi nhận:\n- Như Điều 3;\n- Lưu: VT, Hồ sơ ATTT.")
    doc.save(path)


def _make_checklist(path: Path, ctx: dict) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Checklist ATTT"
    headers = ["STT", "Mã yêu cầu", "Nhóm", "Danh mục", "Yêu cầu", "Hiện trạng", "Trạng thái", "Phương án bổ sung", "Minh chứng", "Người phụ trách", "Hạn hoàn thành"]
    ws.append(headers)
    fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="999999")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
    for idx, answer in enumerate(ctx["answers"], 1):
        req = ctx["requirements"].get(answer.requirement_id)
        ws.append([
            idx,
            _safe(req.code if req else answer.requirement_id),
            _safe(req.group_name if req else ""),
            _safe(req.category if req else ""),
            _safe(req.title if req else ""),
            _safe(answer.current_state),
            _safe(answer.status),
            _safe(answer.improvement_plan),
            _safe(answer.evidence_note) or str(answer.evidence_count or 0),
            _safe(answer.owner),
            answer.due_date.isoformat() if answer.due_date else "",
        ])
    widths = [8, 16, 18, 20, 50, 40, 16, 40, 30, 24, 16]
    for col_idx, width in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(1, col_idx).column_letter].width = width
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
    wb.save(path)


def _make_zip(zip_path: Path, files: list[tuple[str, Path]], evidence: list[EvidenceDocument], include_evidence: bool) -> int:
    evidence_count = 0
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for arcname, path in files:
            zf.write(path, arcname)
        if include_evidence:
            for ev in evidence:
                p = Path(ev.storage_path)
                if p.exists():
                    evidence_count += 1
                    zf.write(p, f"06-MinhChung/{_safe(ev.document_type, 'OTHER')}/{ev.id}_{_slug(ev.original_filename)}")
    return evidence_count


def generate_government_dossier(db: Session, profile_id: int, generated_by: int | None, include_evidence: bool = True, notes: str | None = None) -> GovernmentDossier:
    ctx = _context(db, profile_id)
    profile = ctx["profile"]
    system = ctx["system"]
    max_version = db.scalar(select(func.max(GovernmentDossier.version)).where(GovernmentDossier.profile_id == profile_id)) or 0
    version = int(max_version) + 1
    root = _ensure_dir(DOSSIER_ROOT / f"profile_{profile_id}" / f"v{version}_{uuid4().hex[:10]}")
    base_name = f"HSCD-{_slug(system.code if system else profile.profile_code)}-v{version}"

    generated_files: list[tuple[str, str, str, Path, str, int]] = []
    for file_type, filename, title, sort_order in DOC_TYPES:
        path = root / filename
        if file_type == "TO_TRINH":
            _make_to_trinh(path, ctx)
            content_type = CONTENT_TYPES["docx"]
        elif file_type == "HO_SO_CAP_DO":
            _make_hoso(path, ctx)
            content_type = CONTENT_TYPES["docx"]
        elif file_type == "XIN_Y_KIEN":
            _make_xin_y_kien(path, ctx)
            content_type = CONTENT_TYPES["docx"]
        elif file_type == "QUYET_DINH":
            _make_quyet_dinh(path, ctx)
            content_type = CONTENT_TYPES["docx"]
        else:
            _make_checklist(path, ctx)
            content_type = CONTENT_TYPES["xlsx"]
        generated_files.append((file_type, filename, title, path, content_type, sort_order))

    zip_path = root / f"{base_name}.zip"
    evidence_count = _make_zip(zip_path, [(f[1], f[3]) for f in generated_files], ctx["evidence"], include_evidence)
    package_hash = _sha256(zip_path)

    dossier = GovernmentDossier(
        profile_id=profile_id,
        dossier_code=f"{profile.profile_code}-DOS-v{version}-{uuid4().hex[:6].upper()}",
        title=f"Bộ hồ sơ cấp độ - {_safe(system.name if system else None, profile.profile_code)} - v{version}",
        version=version,
        status="GENERATED",
        package_filename=zip_path.name,
        package_path=str(zip_path),
        package_size=zip_path.stat().st_size,
        checksum_sha256=package_hash,
        included_evidence_count=evidence_count,
        generated_by=generated_by,
        notes=notes,
    )
    db.add(dossier)
    db.flush()

    for file_type, filename, title, path, content_type, sort_order in generated_files:
        db.add(GovernmentDossierFile(
            dossier_id=dossier.id,
            profile_id=profile_id,
            file_type=file_type,
            display_name=title,
            file_name=filename,
            file_path=str(path),
            content_type=content_type,
            file_size=path.stat().st_size,
            checksum_sha256=_sha256(path),
            sort_order=sort_order,
        ))
    db.commit()
    db.refresh(dossier)
    return dossier


def list_dossier_files(db: Session, dossier_id: int) -> list[GovernmentDossierFile]:
    return db.scalars(
        select(GovernmentDossierFile)
        .where(GovernmentDossierFile.dossier_id == dossier_id)
        .order_by(GovernmentDossierFile.sort_order, GovernmentDossierFile.id)
    ).all()


def delete_physical_dossier(dossier: GovernmentDossier) -> None:
    try:
        root = Path(dossier.package_path).parent
        if root.exists() and root.is_dir() and str(root).startswith(str(DOSSIER_ROOT)):
            shutil.rmtree(root, ignore_errors=True)
    except Exception:
        pass
