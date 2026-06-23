from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.document_template import DocumentTemplate
from app.models.evidence_document import EvidenceDocument
from app.models.information_system import InformationSystem
from app.models.level_profile import LevelProfile
from app.models.organization import Organization
from app.models.profile_requirement_answer import ProfileRequirementAnswer
from app.models.security_requirement import SecurityRequirement


DOSSIER_TO_TEMPLATE_TYPE = {
    "TO_TRINH": "APPROVAL_SUBMISSION",
    "HO_SO_CAP_DO": "PROFILE_EXPLANATION",
    "XIN_Y_KIEN": "CONSULTATION_REQUEST",
    "QUYET_DINH": "APPROVAL_DECISION",
    "CHECKLIST": "CHECKLIST_APPENDIX",
}


TEMPLATE_VARIABLES = [
    ("agency_name", "Tên cơ quan/tổ chức ban hành hoặc quản lý hồ sơ", "NGÂN HÀNG CHÍNH SÁCH XÃ HỘI"),
    ("place_name", "Địa danh ban hành văn bản", "Hà Nội"),
    ("document_date", "Ngày sinh văn bản dạng dd/mm/yyyy", "20/06/2026"),
    ("document_datetime", "Thời điểm sinh văn bản", "20/06/2026 08:30"),
    ("profile_id", "ID hồ sơ cấp độ", "3"),
    ("profile_code", "Mã hồ sơ đề xuất cấp độ", "HSCD-CORE-2026"),
    ("profile_name", "Tên hồ sơ", "Hồ sơ đề xuất cấp độ hệ thống Core Banking"),
    ("proposed_level", "Cấp độ đề xuất/phê duyệt", "4"),
    ("profile_status", "Trạng thái hồ sơ", "INTERNALLY_APPROVED"),
    ("basis_for_level", "Căn cứ xác định cấp độ", "Căn cứ dữ liệu, kết nối, CIA..."),
    ("system_scope", "Thuyết minh phạm vi hệ thống", "Phạm vi hệ thống Core Banking..."),
    ("technical_architecture", "Thuyết minh kiến trúc kỹ thuật", "Mô hình 3 lớp, Oracle RAC..."),
    ("confidentiality_impact", "Tác động khi mất tính bí mật", "HIGH"),
    ("integrity_impact", "Tác động khi mất tính toàn vẹn", "HIGH"),
    ("availability_impact", "Tác động khi mất tính sẵn sàng", "CRITICAL"),
    ("system_id", "ID hệ thống thông tin", "1"),
    ("system_code", "Mã hệ thống thông tin", "CORE"),
    ("system_name", "Tên hệ thống thông tin", "Core Banking"),
    ("system_purpose", "Mục tiêu hệ thống", "Quản lý nghiệp vụ tín dụng chính sách"),
    ("system_scope_description", "Phạm vi hệ thống từ danh mục HTTT", "Toàn hệ thống NHCSXH"),
    ("system_functions", "Chức năng chính", "Giao dịch, kế toán, báo cáo..."),
    ("user_groups", "Nhóm người dùng", "Cán bộ nghiệp vụ, quản trị viên"),
    ("data_types", "Loại dữ liệu xử lý", "Dữ liệu khách hàng, tài chính, giao dịch"),
    ("deployment_model", "Mô hình triển khai", "On-premise"),
    ("environment", "Môi trường", "Production"),
    ("importance_level", "Mức độ quan trọng", "CRITICAL"),
    ("system_status", "Trạng thái hoạt động của HTTT", "ACTIVE"),
    ("owner_organization", "Đơn vị chủ quản", "TTCNTT"),
    ("operator_organization", "Đơn vị vận hành", "Phòng QTHT"),
    ("compliance_rate", "Tỷ lệ đáp ứng checklist", "86.50"),
    ("total_requirements", "Tổng số tiêu chí checklist", "120"),
    ("compliant_requirements", "Số tiêu chí đáp ứng", "100"),
    ("not_compliant_requirements", "Số tiêu chí chưa đáp ứng", "15"),
    ("not_applicable_requirements", "Số tiêu chí không áp dụng", "5"),
    ("evidence_count", "Số tài liệu minh chứng", "25"),
    ("signer_title", "Chức danh người ký", "GIÁM ĐỐC"),
    ("signer_name", "Họ tên người ký", "Nguyễn Văn A"),
]


def safe(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def org_name(db: Session, org_id: int | None) -> str:
    if not org_id:
        return "Chưa cập nhật"
    org = db.get(Organization, org_id)
    return safe(org.name if org else None, "Chưa cập nhật")


def get_active_template(db: Session, document_type: str) -> DocumentTemplate | None:
    dtype = DOSSIER_TO_TEMPLATE_TYPE.get(document_type.upper(), document_type.upper())
    candidates = db.scalars(
        select(DocumentTemplate)
        .where(
            DocumentTemplate.document_type == dtype,
            DocumentTemplate.is_active == True,
            DocumentTemplate.template_path.is_not(None),
        )
        .order_by(DocumentTemplate.is_default.desc(), DocumentTemplate.updated_at.desc(), DocumentTemplate.id.desc())
    ).all()
    for item in candidates:
        if item.template_path and Path(item.template_path).exists() and item.file_format.lower() == "docx":
            return item
    return None


def build_template_context(
    db: Session,
    profile_id: int,
    *,
    agency_name: str | None = "NGÂN HÀNG CHÍNH SÁCH XÃ HỘI",
    place_name: str | None = "Hà Nội",
    signer_title: str | None = "THỦ TRƯỞNG ĐƠN VỊ",
    signer_name: str | None = None,
) -> dict[str, Any]:
    profile = db.get(LevelProfile, profile_id)
    if not profile:
        raise ValueError("Level profile not found")
    system = db.get(InformationSystem, profile.information_system_id)

    answers = db.scalars(
        select(ProfileRequirementAnswer)
        .where(ProfileRequirementAnswer.profile_id == profile_id)
        .order_by(ProfileRequirementAnswer.id)
    ).all()
    requirement_ids = [a.requirement_id for a in answers]
    requirements = {}
    if requirement_ids:
        reqs = db.scalars(select(SecurityRequirement).where(SecurityRequirement.id.in_(requirement_ids))).all()
        requirements = {r.id: r for r in reqs}

    evidence = db.scalars(
        select(EvidenceDocument)
        .where(EvidenceDocument.profile_id == profile_id)
        .order_by(EvidenceDocument.created_at.desc(), EvidenceDocument.id.desc())
    ).all()

    compliant = sum(1 for a in answers if safe(a.status).upper() == "COMPLIANT")
    not_compliant = sum(1 for a in answers if safe(a.status).upper() == "NON_COMPLIANT")
    not_applicable = sum(1 for a in answers if safe(a.status).upper() == "NOT_APPLICABLE")
    total_effective = max(1, len(answers) - not_applicable)
    compliance_rate = round(compliant * 100 / total_effective, 2)

    generated_at = datetime.now()
    owner = org_name(db, system.owner_org_id if system else None)
    operator = org_name(db, system.operator_org_id if system else None)

    checklist_rows = []
    for idx, answer in enumerate(answers, 1):
        req = requirements.get(answer.requirement_id)
        checklist_rows.append({
            "stt": idx,
            "requirement_code": safe(req.code if req else answer.requirement_id),
            "requirement_title": safe(req.title if req else "Yêu cầu"),
            "group_name": safe(req.group_name if req else ""),
            "category": safe(req.category if req else ""),
            "current_state": safe(answer.current_state),
            "status": safe(answer.status),
            "improvement_plan": safe(answer.improvement_plan),
            "evidence_note": safe(answer.evidence_note),
            "evidence_count": safe(answer.evidence_count or 0),
            "owner": safe(answer.owner),
            "due_date": answer.due_date.isoformat() if answer.due_date else "",
        })

    evidence_rows = [
        {
            "stt": idx,
            "title": safe(ev.title),
            "document_type": safe(ev.document_type),
            "file_name": safe(ev.original_filename),
            "version": safe(ev.version),
            "uploaded_at": ev.created_at.strftime("%d/%m/%Y %H:%M") if ev.created_at else "",
        }
        for idx, ev in enumerate(evidence, 1)
    ]

    flat = {
        "agency_name": safe(agency_name),
        "place_name": safe(place_name),
        "document_date": generated_at.strftime("%d/%m/%Y"),
        "document_datetime": generated_at.strftime("%d/%m/%Y %H:%M"),
        "profile_id": profile.id,
        "profile_code": safe(profile.profile_code),
        "profile_name": safe(getattr(profile, "name", None), safe(profile.profile_code)),
        "proposed_level": safe(profile.proposed_level),
        "profile_status": safe(profile.status),
        "basis_for_level": safe(profile.basis_for_level),
        "system_scope": safe(profile.system_scope_description),
        "technical_architecture": safe(profile.technical_architecture),
        "confidentiality_impact": safe(profile.confidentiality_impact),
        "integrity_impact": safe(profile.integrity_impact),
        "availability_impact": safe(profile.availability_impact),
        "system_id": system.id if system else "",
        "system_code": safe(system.code if system else ""),
        "system_name": safe(system.name if system else "", safe(profile.profile_code)),
        "system_purpose": safe(system.purpose if system else ""),
        "system_scope_description": safe(system.scope if system else ""),
        "system_functions": safe(system.main_functions if system else ""),
        "user_groups": safe(system.user_groups if system else ""),
        "data_types": safe(system.data_types if system else ""),
        "deployment_model": safe(system.deployment_model if system else ""),
        "environment": safe(system.environment if system else ""),
        "importance_level": safe(system.importance_level if system else ""),
        "system_status": safe(
            getattr(system, "status",
                getattr(system, "operational_status",
                    getattr(system, "system_status", "")
                )
            ) if system else ""
        ),
        "owner_organization": owner,
        "operator_organization": operator,
        "compliance_rate": compliance_rate,
        "total_requirements": len(answers),
        "compliant_requirements": compliant,
        "not_compliant_requirements": not_compliant,
        "not_applicable_requirements": not_applicable,
        "evidence_count": len(evidence),
        "signer_title": safe(signer_title),
        "signer_name": safe(signer_name),
    }

    context: dict[str, Any] = {
        **flat,
        "profile": {
            "id": profile.id,
            "profile_code": flat["profile_code"],
            "name": flat["profile_name"],
            "proposed_level": flat["proposed_level"],
            "status": flat["profile_status"],
            "basis_for_level": flat["basis_for_level"],
            "system_scope_description": flat["system_scope"],
            "technical_architecture": flat["technical_architecture"],
            "confidentiality_impact": flat["confidentiality_impact"],
            "integrity_impact": flat["integrity_impact"],
            "availability_impact": flat["availability_impact"],
        },
        "system": {
            "id": flat["system_id"],
            "code": flat["system_code"],
            "name": flat["system_name"],
            "purpose": flat["system_purpose"],
            "scope": flat["system_scope_description"],
            "main_functions": flat["system_functions"],
            "user_groups": flat["user_groups"],
            "data_types": flat["data_types"],
            "deployment_model": flat["deployment_model"],
            "environment": flat["environment"],
            "importance_level": flat["importance_level"],
            "status": flat["system_status"],
        },
        "organizations": {
            "owner": owner,
            "operator": operator,
        },
        "checklist_rows": checklist_rows,
        "evidence_rows": evidence_rows,
    }
    return context


_PATTERN = re.compile(r"{{\s*([A-Za-z0-9_.]+)\s*}}")


def _lookup(context: dict[str, Any], key: str) -> str:
    if key in context:
        return safe(context[key])
    current: Any = context
    for part in key.split("."):
        if isinstance(current, dict):
            current = current.get(part)
        else:
            current = getattr(current, part, None)
        if current is None:
            return ""
    return safe(current)


def _render_text(text: str, context: dict[str, Any]) -> str:
    return _PATTERN.sub(lambda m: _lookup(context, m.group(1)), text or "")


def _render_paragraph(paragraph, context: dict[str, Any]) -> None:
    original = paragraph.text
    rendered = _render_text(original, context)
    if rendered != original:
        # Replacing paragraph.text keeps the document simple and avoids leaving
        # partially-rendered placeholders split across runs.
        paragraph.text = rendered


def _render_table(table, context: dict[str, Any]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _render_paragraph(paragraph, context)
            for nested in cell.tables:
                _render_table(nested, context)


def render_docx_template(template_path: str | Path, output_path: str | Path, context: dict[str, Any]) -> None:
    doc = Document(str(template_path))
    for paragraph in doc.paragraphs:
        _render_paragraph(paragraph, context)
    for table in doc.tables:
        _render_table(table, context)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _render_paragraph(paragraph, context)
        for table in section.header.tables:
            _render_table(table, context)
        for paragraph in section.footer.paragraphs:
            _render_paragraph(paragraph, context)
        for table in section.footer.tables:
            _render_table(table, context)
    doc.save(str(output_path))


def extract_placeholders_from_docx(template_path: str | Path) -> list[str]:
    doc = Document(str(template_path))
    texts: list[str] = []
    texts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    for section in doc.sections:
        texts.extend(p.text for p in section.header.paragraphs)
        texts.extend(p.text for p in section.footer.paragraphs)
    found = set()
    for text in texts:
        found.update(m.group(1) for m in _PATTERN.finditer(text or ""))
    return sorted(found)


def variable_registry() -> dict[str, Any]:
    return {
        "syntax": "Sử dụng cú pháp {{ variable_name }} hoặc {{ profile.profile_code }} trong file DOCX.",
        "variables": [
            {"key": key, "description": desc, "sample": sample}
            for key, desc, sample in TEMPLATE_VARIABLES
        ],
        "collections": [
            {"key": "checklist_rows", "description": "Danh sách checklist dạng dữ liệu để phát triển template nâng cao."},
            {"key": "evidence_rows", "description": "Danh sách tài liệu minh chứng dạng dữ liệu để phát triển template nâng cao."},
        ],
    }
