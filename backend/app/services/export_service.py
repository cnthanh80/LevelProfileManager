from __future__ import annotations

import hashlib
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Pt
from fastapi import HTTPException
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from app.services.pdf_font import apply_unicode_styles, register_unicode_fonts
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.exported_document import ExportedDocument
from app.models.information_system import InformationSystem
from app.models.level_profile import LevelProfile
from app.models.profile_requirement_answer import ProfileRequirementAnswer
from app.models.security_requirement import SecurityRequirement
from app.models.evidence_document import EvidenceDocument

EXPORT_ROOT = Path("/app/storage/exports")

DOC_TYPE_TITLES = {
    "PROFILE_EXPLANATION": "Thuyết minh hồ sơ đề xuất cấp độ",
    "CHECKLIST_APPENDIX": "Phụ lục checklist đáp ứng yêu cầu ATTT",
    "APPROVAL_SUBMISSION": "Tờ trình phê duyệt hồ sơ đề xuất cấp độ",
    "CONSULTATION_REQUEST": "Văn bản xin ý kiến chuyên môn",
}


def _safe(value) -> str:
    if value is None:
        return ""
    return str(value)


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _get_profile(db: Session, profile_id: int) -> tuple[LevelProfile, InformationSystem]:
    profile = db.get(LevelProfile, profile_id)
    if not profile:
        raise HTTPException(status_code=404, detail="Level profile not found")
    system = db.get(InformationSystem, profile.information_system_id)
    if not system:
        raise HTTPException(status_code=404, detail="Information system not found")
    return profile, system


def _get_checklist_rows(db: Session, profile_id: int):
    stmt = (
        select(ProfileRequirementAnswer, SecurityRequirement)
        .join(SecurityRequirement, ProfileRequirementAnswer.requirement_id == SecurityRequirement.id)
        .where(ProfileRequirementAnswer.profile_id == profile_id)
        .order_by(SecurityRequirement.group_name, SecurityRequirement.category, SecurityRequirement.code)
    )
    return db.execute(stmt).all()


def _get_evidence_count(db: Session, profile_id: int) -> int:
    return len(db.scalars(select(EvidenceDocument).where(EvidenceDocument.profile_id == profile_id)).all())


def _add_doc_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Thông tin"
    hdr[1].text = "Nội dung"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value or ""
    return table


def generate_docx(db: Session, profile_id: int, document_type: str, generated_by: int | None) -> ExportedDocument:
    profile, system = _get_profile(db, profile_id)
    checklist_rows = _get_checklist_rows(db, profile_id)
    evidence_count = _get_evidence_count(db, profile_id)

    EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    title = DOC_TYPE_TITLES.get(document_type, document_type)
    stored_filename = f"export_{profile_id}_{document_type.lower()}_{uuid4().hex}.docx"
    output_path = EXPORT_ROOT / stored_filename

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    _add_doc_heading(doc, title.upper(), 0)
    doc.add_paragraph(f"Mã hồ sơ: {profile.profile_code}")
    doc.add_paragraph(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    _add_doc_heading(doc, "1. Thông tin hệ thống thông tin", 1)
    _add_key_value_table(doc, [
        ("Tên hệ thống", _safe(system.name)),
        ("Mã hệ thống", _safe(system.code)),
        ("Mục tiêu", _safe(system.purpose)),
        ("Phạm vi", _safe(system.scope)),
        ("Chức năng chính", _safe(system.main_functions)),
        ("Nhóm người dùng", _safe(system.user_groups)),
        ("Loại dữ liệu xử lý", _safe(system.data_types)),
        ("Mô hình triển khai", _safe(system.deployment_model)),
        ("Môi trường", _safe(system.environment)),
        ("Trạng thái hoạt động", _safe(system.operation_status)),
    ])

    _add_doc_heading(doc, "2. Thông tin hồ sơ đề xuất cấp độ", 1)
    _add_key_value_table(doc, [
        ("Cấp độ đề xuất", str(profile.proposed_level)),
        ("Trạng thái", _safe(profile.status)),
        ("Căn cứ xác định cấp độ", _safe(profile.basis_for_level)),
        ("Thuyết minh phạm vi", _safe(profile.system_scope_description)),
        ("Kiến trúc kỹ thuật", _safe(profile.technical_architecture)),
        ("Tác động mất bí mật", _safe(profile.confidentiality_impact)),
        ("Tác động mất toàn vẹn", _safe(profile.integrity_impact)),
        ("Tác động mất sẵn sàng", _safe(profile.availability_impact)),
        ("Số tài liệu minh chứng", str(evidence_count)),
    ])

    _add_doc_heading(doc, "3. Phụ lục checklist yêu cầu ATTT", 1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Mã", "Nhóm", "Danh mục", "Yêu cầu", "Trạng thái", "Hiện trạng", "Minh chứng"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for answer, req in checklist_rows:
        cells = table.add_row().cells
        cells[0].text = _safe(req.code)
        cells[1].text = _safe(req.group_name)
        cells[2].text = _safe(req.category)
        cells[3].text = _safe(req.title)
        cells[4].text = _safe(answer.status)
        cells[5].text = _safe(answer.current_state)
        cells[6].text = str(answer.evidence_count or 0)

    doc.add_page_break()
    doc.add_paragraph("Ghi chú: Tài liệu được xuất tự động từ hệ thống Level Profile Manager.")
    doc.save(output_path)

    return _record_export(db, profile_id, document_type, "docx", title, output_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", generated_by)


def generate_pdf(db: Session, profile_id: int, document_type: str, generated_by: int | None) -> ExportedDocument:
    profile, system = _get_profile(db, profile_id)
    checklist_rows = _get_checklist_rows(db, profile_id)
    evidence_count = _get_evidence_count(db, profile_id)

    EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    title = DOC_TYPE_TITLES.get(document_type, document_type)
    stored_filename = f"export_{profile_id}_{document_type.lower()}_{uuid4().hex}.pdf"
    output_path = EXPORT_ROOT / stored_filename

    styles = getSampleStyleSheet()
    regular_font, bold_font = apply_unicode_styles(styles)
    story = []
    story.append(Paragraph(title.upper(), styles["Title"]))
    story.append(Paragraph(f"Mã hồ sơ: {profile.profile_code}", styles["Normal"]))
    story.append(Paragraph(f"Ngày xuất: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("1. Thông tin hệ thống thông tin", styles["Heading2"]))
    profile_table = Table([
        ["Tên hệ thống", _safe(system.name)],
        ["Mã hệ thống", _safe(system.code)],
        ["Cấp độ đề xuất", str(profile.proposed_level)],
        ["Trạng thái", _safe(profile.status)],
        ["Mô hình triển khai", _safe(system.deployment_model)],
        ["Môi trường", _safe(system.environment)],
        ["Số tài liệu minh chứng", str(evidence_count)],
    ], colWidths=[140, 350])
    profile_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), regular_font),
    ]))
    story.append(profile_table)
    story.append(Spacer(1, 12))
    story.append(Paragraph("2. Checklist tóm tắt", styles["Heading2"]))
    data = [["Mã", "Nhóm", "Yêu cầu", "Trạng thái", "Minh chứng"]]
    for answer, req in checklist_rows[:80]:
        data.append([_safe(req.code), _safe(req.group_name), _safe(req.title)[:80], _safe(answer.status), str(answer.evidence_count or 0)])
    table = Table(data, colWidths=[55, 80, 230, 80, 60])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), regular_font),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("FONTNAME", (0, 0), (-1, -1), regular_font),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph("Ghi chú: Bản PDF được sinh tự động từ hệ thống Level Profile Manager.", styles["Normal"]))

    pdf = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    pdf.build(story)

    return _record_export(db, profile_id, document_type, "pdf", title, output_path, "application/pdf", generated_by)


def _record_export(db: Session, profile_id: int, document_type: str, file_format: str, title: str, output_path: Path, content_type: str, generated_by: int | None) -> ExportedDocument:
    checksum = _sha256(output_path)
    item = ExportedDocument(
        profile_id=profile_id,
        document_type=document_type,
        file_format=file_format,
        title=title,
        original_filename=output_path.name,
        stored_filename=output_path.name,
        storage_path=str(output_path),
        content_type=content_type,
        file_size=output_path.stat().st_size,
        checksum_sha256=checksum,
        generated_by=generated_by,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


def export_profile_document(db: Session, profile_id: int, document_type: str, file_format: str, generated_by: int | None) -> ExportedDocument:
    document_type = document_type.upper()
    file_format = file_format.lower()
    if document_type not in DOC_TYPE_TITLES:
        raise HTTPException(status_code=400, detail=f"Unsupported document_type. Allowed: {', '.join(DOC_TYPE_TITLES)}")
    if file_format == "docx":
        return generate_docx(db, profile_id, document_type, generated_by)
    if file_format == "pdf":
        return generate_pdf(db, profile_id, document_type, generated_by)
    raise HTTPException(status_code=400, detail="Unsupported file_format. Allowed: docx, pdf")
